"""
Layer 1 + 2 — CV / Resume Parser

Accepts: PDF, DOCX, plain text
Extracts: name, contact info, work experience, education, skills,
          certifications, languages, total years of experience.
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Union

import spacy

from cv_matcher.models import Education, ParsedCV, WorkExperience
from cv_matcher.llm.client import get_client
from cv_matcher.llm.prompts import CV_PARSE_SYSTEM, CV_PARSE_USER

# ── NLP model (loaded once) ───────────────────────────────────────────────────
try:
    _nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess, sys
    subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
    _nlp = spacy.load("en_core_web_sm")


# ── Section header patterns ───────────────────────────────────────────────────
_SECTION_PATTERNS = {
    "skills":         re.compile(r"^(technical\s+)?skills?(\s+&\s+\w+)?", re.I),
    "experience":     re.compile(r"^(work\s+|professional\s+)?experience|employment|history", re.I),
    "education":      re.compile(r"^education|academic|qualifications?", re.I),
    "certifications": re.compile(r"^certif|licenses?|accreditations?", re.I),
    "languages":      re.compile(r"^languages?(\s+spoken)?", re.I),
    "summary":        re.compile(r"^(professional\s+)?(summary|profile|objective|about)", re.I),
}

_SENIORITY_KEYWORDS = {
    "junior":    re.compile(r"\bjunior|jr\.?\b|entry.?level|intern\b", re.I),
    "mid":       re.compile(r"\bmid.?level|associate\b", re.I),
    "senior":    re.compile(r"\bsenior|sr\.?\b|principal\b", re.I),
    "lead":      re.compile(r"\blead|staff|head\s+of|director|manager|architect\b", re.I),
}

_TECH_SKILLS_VOCAB = {
    # Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "bash", "shell",
    # Web
    "react", "react.js", "angular", "vue", "next.js", "nuxt", "svelte",
    "html", "css", "sass", "tailwind", "bootstrap", "graphql", "rest", "soap",
    # Backend / frameworks
    "node.js", "express", "django", "flask", "fastapi", "spring", "rails",
    "laravel", ".net", "asp.net",
    # Cloud / DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "k8s", "terraform",
    "ansible", "jenkins", "github actions", "ci/cd", "helm", "prometheus",
    "grafana", "linux", "nginx", "apache",
    # Data / ML
    "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "kafka", "spark", "hadoop", "airflow", "dbt", "pandas", "numpy",
    "scikit-learn", "tensorflow", "pytorch", "keras", "mlflow", "langchain",
    "hugging face", "openai", "llm",
    # Soft / process
    "agile", "scrum", "jira", "confluence", "git", "github", "gitlab",
}

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", re.I)
_PHONE_RE = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w-]+", re.I)
_DURATION_RE = re.compile(
    r"(\d+)\s*\+?\s*years?(?:\s+(?:of\s+)?experience)?", re.I
)


class CVParser:
    """Parse a CV/resume from PDF, DOCX, or plain text."""

    # ── Public entry point ────────────────────────────────────────────────────

    def parse(self, source: Union[str, Path, bytes], enrich: bool = True) -> ParsedCV:
        """
        Parse *source* into a structured ParsedCV.
        If enrich=True and ANTHROPIC_API_KEY is set, LLM fills in any fields
        that regex/NLP couldn't reliably extract.
        """
        text = self._load_text(source)
        cv = self._parse_text(text)
        if enrich:
            cv = self._llm_enrich(cv, text)
        return cv

    def _llm_enrich(self, cv: ParsedCV, raw_text: str) -> ParsedCV:
        """Use LLM to fill gaps left by regex/NLP parsing."""
        llm = get_client()
        if not llm.available:
            return cv

        user = CV_PARSE_USER.format(cv_text=raw_text[:6000])  # truncate for token safety
        data = llm.complete_json(CV_PARSE_SYSTEM, user, max_tokens=800)
        if not data or not isinstance(data, dict):
            return cv

        # Only overwrite fields that regex left empty or weak
        if not cv.name and data.get("name"):
            cv.name = data["name"]
        if not cv.email and data.get("email"):
            cv.email = data["email"]
        if not cv.phone and data.get("phone"):
            cv.phone = data["phone"]
        if not cv.linkedin and data.get("linkedin"):
            cv.linkedin = data["linkedin"]
        if not cv.total_experience_years and data.get("total_experience_years"):
            cv.total_experience_years = float(data["total_experience_years"])
        if not cv.languages and data.get("languages"):
            cv.languages = data["languages"]
        if not cv.certifications and data.get("certifications"):
            cv.certifications = data["certifications"]

        # Merge LLM skills with regex-extracted ones (union, deduplicated)
        llm_skills = [s.lower() for s in data.get("skills", [])]
        merged = list(dict.fromkeys(cv.skills + llm_skills))
        cv.skills = merged

        return cv

    # ── Text loading ──────────────────────────────────────────────────────────

    def _load_text(self, source: Union[str, Path, bytes]) -> str:
        if isinstance(source, bytes):
            # Try PDF first, then DOCX
            return self._bytes_to_text(source)
        path = Path(source)
        if path.exists():
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                return self._pdf_to_text(path)
            elif suffix in (".docx", ".doc"):
                return self._docx_to_text(path)
            else:
                return path.read_text(encoding="utf-8", errors="replace")
        # Treat as raw text
        return str(source)

    def _pdf_to_text(self, path: Path) -> str:
        import pdfplumber
        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return "\n".join(pages)

    def _docx_to_text(self, path: Path) -> str:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _bytes_to_text(self, data: bytes) -> str:
        # Detect magic bytes
        if data[:4] == b"%PDF":
            import pdfplumber
            pages: list[str] = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            return "\n".join(pages)
        elif data[:2] == b"PK":
            from docx import Document
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return data.decode("utf-8", errors="replace")

    # ── Core parsing ──────────────────────────────────────────────────────────

    def _parse_text(self, text: str) -> ParsedCV:
        cv = ParsedCV(raw_text=text)

        # Contact info (quick regex pass before NLP)
        cv.email = self._extract_email(text)
        cv.phone = self._extract_phone(text)
        cv.linkedin = self._extract_linkedin(text)

        # Segment text into sections
        sections = self._segment_sections(text)

        # Name: try NLP on first 10 lines
        cv.name = self._extract_name(text)

        # Skills
        cv.skills = self._extract_skills(
            sections.get("skills", "") + "\n" + sections.get("summary", "") + "\n" + text
        )

        # Experience
        cv.work_experience = self._extract_experience(sections.get("experience", ""))
        cv.total_experience_years = self._compute_total_experience(
            cv.work_experience, text
        )

        # Education
        cv.education = self._extract_education(sections.get("education", ""))

        # Certifications
        cv.certifications = self._extract_list_items(sections.get("certifications", ""))

        # Languages
        cv.languages = self._extract_languages(sections.get("languages", ""))

        return cv

    # ── Section segmentation ──────────────────────────────────────────────────

    def _segment_sections(self, text: str) -> dict[str, str]:
        lines = text.splitlines()
        sections: dict[str, list[str]] = {"_default": []}
        current = "_default"

        for line in lines:
            stripped = line.strip()
            matched_section = self._classify_section_header(stripped)
            if matched_section and len(stripped) < 60:
                current = matched_section
                sections.setdefault(current, [])
            else:
                sections.setdefault(current, []).append(line)

        return {k: "\n".join(v) for k, v in sections.items()}

    def _classify_section_header(self, line: str) -> str | None:
        for name, pat in _SECTION_PATTERNS.items():
            if pat.match(line.strip(":–—●•").strip()):
                return name
        return None

    # ── Field extractors ──────────────────────────────────────────────────────

    def _extract_email(self, text: str) -> str:
        m = _EMAIL_RE.search(text)
        return m.group(0) if m else ""

    def _extract_phone(self, text: str) -> str:
        m = _PHONE_RE.search(text)
        return m.group(0).strip() if m else ""

    def _extract_linkedin(self, text: str) -> str:
        m = _LINKEDIN_RE.search(text)
        return m.group(0) if m else ""

    def _extract_name(self, text: str) -> str:
        # Use spaCy PERSON entities from the first 500 chars
        snippet = text[:500]
        doc = _nlp(snippet)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                return ent.text.strip()
        # Fallback: first non-empty, non-contact line
        for line in text.splitlines():
            line = line.strip()
            if line and not _EMAIL_RE.search(line) and not _PHONE_RE.search(line):
                if len(line.split()) <= 5:
                    return line
        return ""

    def _extract_skills(self, text: str) -> list[str]:
        found: set[str] = set()
        lower = text.lower()

        # Match against vocabulary
        for skill in _TECH_SKILLS_VOCAB:
            pattern = r"\b" + re.escape(skill) + r"\b"
            if re.search(pattern, lower):
                found.add(skill)

        # Also pull comma/bullet separated items from skill sections
        for token in re.split(r"[,\n•\|/]", text):
            token = token.strip(" :-–—●•\t")
            if 2 <= len(token) <= 40:
                tl = token.lower()
                if tl in _TECH_SKILLS_VOCAB:
                    found.add(tl)

        return sorted(found)

    def _extract_experience(self, text: str) -> list[WorkExperience]:
        experiences: list[WorkExperience] = []
        if not text.strip():
            return experiences

        # Split on likely job-entry boundaries (blank lines or date-like patterns)
        blocks = re.split(r"\n{2,}|\n(?=\d{4})", text)
        for block in blocks:
            block = block.strip()
            if not block or len(block) < 20:
                continue
            exp = WorkExperience()
            lines = block.splitlines()

            # Title is usually the first substantial line
            if lines:
                exp.title = lines[0].strip()

            # Company: look for patterns like "@ Company" or second line
            for line in lines[1:3]:
                if re.search(r"\bat\b|@|\|", line, re.I):
                    exp.company = re.split(r"\bat\b|@|\|", line, maxsplit=1)[-1].strip()
                    break

            # Duration: count months from date ranges
            exp.duration_months = self._extract_duration_months(block)

            # Skills mentioned inline
            exp.skills_mentioned = self._extract_skills(block)
            exp.description = block

            if exp.title:
                experiences.append(exp)

        return experiences

    def _extract_duration_months(self, text: str) -> int:
        # Pattern: "Jan 2020 – Mar 2022"
        date_re = re.compile(
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*"
            r"[\s,]+(\d{4})\s*[-–—to]+\s*"
            r"(?:(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*"
            r"[\s,]+(\d{4})|[Pp]resent|[Cc]urrent|[Nn]ow)",
            re.I,
        )
        months_map = {
            "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
            "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
        }
        import datetime
        total = 0
        for m in date_re.finditer(text):
            start_m = months_map.get(m.group(1)[:3].lower(), 1)
            start_y = int(m.group(2))
            if m.group(4):
                end_m = months_map.get(m.group(3)[:3].lower(), 1)
                end_y = int(m.group(4))
            else:
                now = datetime.date.today()
                end_m, end_y = now.month, now.year
            total += (end_y - start_y) * 12 + (end_m - start_m)
        return max(total, 0)

    def _compute_total_experience(
        self, experiences: list[WorkExperience], raw_text: str
    ) -> float:
        months = sum(e.duration_months for e in experiences)
        if months > 0:
            return round(months / 12, 1)
        # Fallback: regex like "5+ years of experience"
        m = _DURATION_RE.search(raw_text)
        if m:
            return float(m.group(1))
        return 0.0

    def _extract_education(self, text: str) -> list[Education]:
        educations: list[Education] = []
        if not text.strip():
            return educations

        degree_re = re.compile(
            r"(B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|Bachelor|Master|Doctor|MBA|BSc|MSc)",
            re.I,
        )
        year_re = re.compile(r"\b(19|20)\d{2}\b")

        blocks = re.split(r"\n{2,}", text)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            edu = Education()
            dm = degree_re.search(block)
            if dm:
                edu.degree = dm.group(0)
            ym = year_re.search(block)
            if ym:
                edu.year = int(ym.group(0))
            # Institution: look for "University", "College", "Institute"
            inst_m = re.search(
                r"([\w\s]+(?:University|College|Institute|School|Academy)[\w\s]*)",
                block, re.I,
            )
            if inst_m:
                edu.institution = inst_m.group(0).strip()
            if edu.degree or edu.institution:
                educations.append(edu)

        return educations

    def _extract_list_items(self, text: str) -> list[str]:
        items: list[str] = []
        for line in text.splitlines():
            line = line.strip(" :-–—●•\t")
            if line and len(line) > 3:
                items.append(line)
        return items

    def _extract_languages(self, text: str) -> list[str]:
        known = {
            "english", "french", "arabic", "spanish", "german", "chinese",
            "japanese", "portuguese", "italian", "russian", "dutch", "korean",
            "hindi", "turkish", "polish", "swedish", "norwegian",
        }
        found: list[str] = []
        lower = text.lower()
        for lang in known:
            if re.search(r"\b" + lang + r"\b", lower):
                found.append(lang.capitalize())
        # Also capture bullet-listed items
        for item in self._extract_list_items(text):
            if item.lower() not in {l.lower() for l in found}:
                found.append(item)
        return list(dict.fromkeys(found))  # deduplicate, preserve order
